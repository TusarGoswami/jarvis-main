import io
import json
import re
from typing import Dict, Any, List, Optional
from pydantic import BaseModel
from google import genai
from google.genai import types

from app.config import settings
from app.core.llm_provider import generate_gemini_content_sync

# Supported technical domains
TECHNICAL_DOMAINS = [
    "Full Stack Development",
    "Frontend Development",
    "Backend Development",
    "Software Engineering",
    "Data Science",
    "Machine Learning",
    "AI / ML",
    "DevOps",
    "Cloud Engineering",
    "Cybersecurity",
    "Mobile Development",
    "Java Development",
    "Python Development",
    "C++ Development",
    "JavaScript Development",
    "Database / SQL",
    "System Design",
]

def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Extracts raw UTF-8 text from PDF, DOCX, or plain text files.
    """
    if not file_bytes:
        raise ValueError("Uploaded file is empty.")
    
    fname_lower = filename.lower()
    
    # 1. Plain text / Markdown
    if fname_lower.endswith((".txt", ".md", ".log", ".rtf")):
        try:
            return file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            return file_bytes.decode("latin-1", errors="ignore")
            
    # 2. PDF Extraction
    if fname_lower.endswith(".pdf"):
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            pages_text = []
            for page in reader.pages:
                txt = page.extract_text()
                if txt:
                    pages_text.append(txt)
            combined = "\n".join(pages_text).strip()
            if not combined:
                raise ValueError("PDF contains no readable text or is image-based.")
            return combined
        except Exception as e:
            raise ValueError(f"Failed to parse PDF: {str(e)}")
            
    # 3. DOCX Extraction
    if fname_lower.endswith((".docx", ".doc")):
        try:
            import docx
            doc = docx.Document(io.BytesIO(file_bytes))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    row_txt = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_txt:
                        paragraphs.append(row_txt)
            combined = "\n".join(paragraphs).strip()
            if not combined:
                raise ValueError("DOCX document contains no text.")
            return combined
        except Exception as e:
            raise ValueError(f"Failed to parse DOCX: {str(e)}")

    # Fallback to UTF-8 decoding
    try:
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception:
        raise ValueError(f"Unsupported file format: {filename}")



def parse_resume_with_ai(raw_text: str) -> Dict[str, Any]:
    """
    Extracts structured candidate profile from resume text using Gemini 2.5 Flash
    with fallback parsing for offline environments.
    """
    if settings.GEMINI_API_KEYS and len(raw_text.strip()) > 30:
        prompt = f"""
        You are a high-precision Technical Recruiter AI.
        Analyze the following candidate Resume/CV text and extract structured profile data in strict JSON format.

        Resume Text:
        \"\"\"
        {raw_text[:8000]}
        \"\"\"

        Respond ONLY with a valid JSON object matching this schema:
        {{
            "name": "Candidate Full Name (or 'Candidate' if missing)",
            "email": "Email address or null",
            "phone": "Phone number or null",
            "education": ["Degree/University entries"],
            "skills": ["List of core technical and domain skills"],
            "languages": ["Programming languages known, e.g. Python, Java, TypeScript"],
            "frameworks": ["Frameworks and tools, e.g. React, Next.js, FastAPI, Docker"],
            "projects": [
                {{
                    "title": "Project Title",
                    "description": "Short 1-line description of project and technologies"
                }}
            ],
            "experience_years": "Estimated total years of experience, e.g., '2+ years' or 'Fresher'",
            "summary": "2-sentence executive summary of the candidate's technical profile"
        }}
        """
        try:
            response = generate_gemini_content_sync(
                contents=prompt,
                config=types.GenerateContentConfig(
                    temperature=0.1,
                    response_mime_type="application/json",
                )
            )
            parsed = json.loads(response.text)
            return parsed
        except Exception as e:
            print(f"[InterviewExtractor] AI Resume extraction error: {e}, using heuristic fallback.")

    # Offline / Heuristic Fallback
    return _heuristic_resume_fallback(raw_text)

def parse_job_description_with_ai(raw_text: str) -> Dict[str, Any]:
    """
    Extracts structured job specifications and detects technical domain from JD.
    """
    if settings.GEMINI_API_KEYS and len(raw_text.strip()) > 20:
        prompt = f"""
        You are an AI Technical Hiring Specialist.
        Analyze this Job Description and extract key requirements in strict JSON format.
        Also select the closest matching domain from this exact list:
        {json.dumps(TECHNICAL_DOMAINS)}

        Job Description Text:
        \"\"\"
        {raw_text[:8000]}
        \"\"\"

        Respond ONLY with a valid JSON object matching this schema:
        {{
            "title": "Job Role / Title (e.g. Senior Backend Engineer)",
            "company": "Company Name (or 'Confidential/Unknown')",
            "required_skills": ["Essential technical requirements"],
            "preferred_skills": ["Nice to have / optional skills"],
            "languages": ["Required programming languages"],
            "frameworks": ["Required frameworks / libraries / cloud tools"],
            "responsibilities": ["Top 3-4 key job responsibilities"],
            "experience_required": "Required years/level (e.g. '3-5 Years', 'Fresher', 'Senior')",
            "inferred_domain": "One exact match from the provided domain list",
            "keywords": ["Top 5-8 searchable technical keywords"]
        }}
        """
        try:
            response = generate_gemini_content_sync(
                contents=prompt,
                config=types.GenerateContentConfig(
                    temperature=0.1,
                    response_mime_type="application/json",
                )
            )
            parsed = json.loads(response.text)
            return parsed
        except Exception as e:
            print(f"[InterviewExtractor] AI JD extraction error: {e}, using heuristic fallback.")

    # Offline / Heuristic Fallback
    return _heuristic_jd_fallback(raw_text)

def _heuristic_resume_fallback(text: str) -> Dict[str, Any]:
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    name = lines[0] if lines else "Candidate Profile"
    
    # Common tech keywords
    tech_keywords = [
        "python", "javascript", "typescript", "java", "c++", "c#", "go", "rust", "sql",
        "react", "next.js", "node.js", "fastapi", "django", "flask", "docker", "kubernetes",
        "aws", "gcp", "azure", "graphql", "mongodb", "postgresql", "git", "linux"
    ]
    
    found_skills = []
    text_lower = text.lower()
    for kw in tech_keywords:
        if re.search(rf"\b{re.escape(kw)}\b", text_lower):
            found_skills.append(kw.title() if kw not in ["sql", "aws", "gcp"] else kw.upper())
            
    return {
        "name": name[:50],
        "email": None,
        "phone": None,
        "education": ["Computer Science / Engineering Background"],
        "skills": found_skills[:12] or ["Software Engineering", "Problem Solving", "System Architecture"],
        "languages": [s for s in found_skills if s.lower() in ["python", "javascript", "typescript", "java", "c++", "go", "rust"]],
        "frameworks": [s for s in found_skills if s.lower() in ["react", "next.js", "fastapi", "django", "docker", "node.js"]],
        "projects": [
            {"title": "Technical Application Portfolio", "description": "Demonstrated software engineering and development projects."}
        ],
        "experience_years": "1-3 Years",
        "summary": f"Technical candidate with demonstrated skills in {', '.join(found_skills[:4]) if found_skills else 'Software Engineering'}."
    }

def _heuristic_jd_fallback(text: str) -> Dict[str, Any]:
    text_lower = text.lower()
    inferred_domain = "Software Engineering"
    
    domain_map = {
        "frontend": "Frontend Development",
        "react": "Frontend Development",
        "backend": "Backend Development",
        "full stack": "Full Stack Development",
        "fullstack": "Full Stack Development",
        "machine learning": "Machine Learning",
        "ai": "AI / ML",
        "data science": "Data Science",
        "devops": "DevOps",
        "cloud": "Cloud Engineering",
        "security": "Cybersecurity",
        "mobile": "Mobile Development",
        "android": "Mobile Development",
        "ios": "Mobile Development",
        "system design": "System Design",
    }
    
    for key, dom in domain_map.items():
        if key in text_lower:
            inferred_domain = dom
            break

    return {
        "title": "Software Engineer",
        "company": "Technology Organization",
        "required_skills": ["Software Engineering", "Core Fundamentals", "Data Structures", "System Design"],
        "preferred_skills": ["Cloud Infrastructure", "API Design", "Distributed Systems"],
        "languages": ["Python", "TypeScript", "Java"],
        "frameworks": ["React", "FastAPI", "Docker"],
        "responsibilities": ["Design and develop robust applications", "Collaborate on architecture and technical solutions"],
        "experience_required": "1–3 Years",
        "inferred_domain": inferred_domain,
        "keywords": ["Architecture", "Engineering", "Algorithms", "Development", "Scalability"]
    }
